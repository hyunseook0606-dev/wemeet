# -*- coding: utf-8 -*-
"""플랫폼 전반 흐름도 DOCX 생성 — 발표용."""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
except ImportError:
    print('python-docx 설치 필요: pip install python-docx')
    sys.exit(1)

OUT = Path(__file__).parent.parent / '위밋모빌리티_플랫폼_흐름도.docx'
doc = Document()

s = doc.sections[0]
s.left_margin = s.right_margin = Cm(2.0)
s.top_margin  = s.bottom_margin = Cm(1.8)

# ── 색상 팔레트 ──────────────────────────────────────────────
C_DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
C_MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
C_LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
C_GREEN      = RGBColor(0x1E, 0x88, 0x50)
C_ORANGE     = RGBColor(0xE6, 0x5C, 0x00)
C_RED        = RGBColor(0xC0, 0x00, 0x00)
C_GRAY       = RGBColor(0x60, 0x60, 0x60)
C_YELLOW_BG  = 'FFF9C4'
C_GREEN_BG   = 'D4EDDA'
C_BLUE_BG    = 'D6E4F0'
C_ORANGE_BG  = 'FFE0B2'
C_RED_BG     = 'FDECEA'

def _shd(p, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    p._p.get_or_add_pPr().append(shd)

def _cell_shd(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def h1(text, color=C_DARK_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(5)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(15)
    r.font.color.rgb = color

def h2(text, color=C_MID_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = color

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10.5)
    r.font.color.rgb = C_GRAY

def body(text, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(size)

def bullet(text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10)

def box(text, fill_hex, text_color=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    _shd(p, fill_hex)
    r = p.add_run(text)
    r.font.size = Pt(10)
    if text_color:
        r.font.color.rgb = text_color

def note(text):
    box(text, 'FFF9C4', RGBColor(0x70, 0x50, 0x00))

def arrow_row(items):
    """items: list of (text, fill_hex) — 화살표로 연결된 흐름 박스"""
    cols = len(items) * 2 - 1
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (txt, fill) in enumerate(items):
        ci = i * 2
        cell = table.cell(0, ci)
        cell.width = Cm(3.5)
        _cell_shd(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.size = Pt(9.5)
        r.bold = True
        if i < len(items) - 1:
            arrow_cell = table.cell(0, ci + 1)
            arrow_cell.width = Cm(0.6)
            ap = arrow_cell.paragraphs[0]
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.add_run('->').font.size = Pt(11)
    doc.add_paragraph()

def add_table(headers, rows, col_widths=None, header_fill='1F4E79'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _cell_shd(hdr_cells[i], header_fill)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        fill = 'F7FBFF' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            _cell_shd(cells[c_idx], fill)
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(30)
tr = title_p.add_run('해상 리스크 대응형 공동 물류 운영·국내운송 연계 플랫폼')
tr.bold = True; tr.font.size = Pt(18)
tr.font.color.rgb = C_DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(4)
sub.add_run('플랫폼 전체 흐름도 및 구조 설명서').font.size = Pt(13)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_after = Pt(4)
sub2.add_run('위밋모빌리티(Wemet Mobility) x 한국해양수산개발원(KMI)').font.size = Pt(11)

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub3.paragraph_format.space_after = Pt(30)
sub3.add_run('2026학년도 1학기 기업문제해결 프로젝트 | 2026.04.29 ~ 2026.05.20').font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. 플랫폼 개요
# ════════════════════════════════════════════════════════════
h1('1. 플랫폼 개요 및 배경')

h2('1-1. 해결하는 문제')
box('중소 수출기업은 홍해 사태·태풍·항만 파업 등 해상 리스크 발생 시 대응 수단이 없어 납기 지연·비용 급증을 그대로 감수해야 합니다. 기존 플랫폼(트레드링스, Flexport)은 가시성·ETA 예측에 강하지만 선제적 의사결정 지원은 부재합니다.', 'FDECEA')

h2('1-2. 플랫폼이 제공하는 것')
add_table(
    ['영역', '기존 플랫폼', '본 플랫폼'],
    [
        ['해상 가시성 / ETA', '강점 (트레드링스, Flexport)', '보완재 포지셔닝'],
        ['선제적 리스크 의사결정', '없음', 'MRI 5차원 AHP + 시나리오 자동 분류'],
        ['항만 창고 자동 탐색', '없음', 'NLIC 439개 + 카카오 실데이터'],
        ['국내 내륙 배차 통합', '없음', '루티(ROOUTY) JSON 자동 생성'],
        ['과거 유사사례 인사이트', '없음', 'historical_matcher (7개 실제 사건 DB)'],
    ],
    [4, 5, 6],
    '1F4E79',
)

h2('1-3. 연계 기업 및 핵심 KPI')
add_table(
    ['지표', '값', '근거'],
    [
        ['LSTM 물동량 예측 MAPE', '9.4%', '시간순 80/20 분할, 원단위(TEU) 역정규화'],
        ['CO2 절감률', '35%', '적재율 55% → 85% 향상 (CBAM 대응)'],
        ['지정학 시나리오 지연', '+14일', '케이프타운 우회 실제 소요일'],
        ['지정학 시나리오 운임', '+30%', '홍해 사태 실제 증가율 (UNCTAD 2024)'],
        ['수에즈 통항 감소', '42~90%', 'UNCTAD 2024 실측'],
        ['AHP 일관성 비율', 'CR=3.1%', '< 10% 기준 통과'],
    ],
    [5, 3, 7],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. 전체 흐름도 (4단계)
# ════════════════════════════════════════════════════════════
h1('2. 플랫폼 전체 흐름도 (4단계)')

body('플랫폼은 화주 입력부터 루티 JSON 생성까지 4단계로 구성됩니다. MRI 임계값이 핵심 분기점입니다.', 10)

arrow_row([
    ('Step 1\n화주 입력', 'D6E4F0'),
    ('Step 2\nMRI 분석', 'FFE0B2'),
    ('Step 3\n창고 추천\n(MRI >= 0.5)', 'D4EDDA'),
    ('Step 4\n루티 JSON', 'EDE7F6'),
])

# Step 1
h2('Step 1 — 화주 입력')
box('화주가 화물 정보를 등록합니다. 이후 전 단계의 기준이 됩니다.', 'D6E4F0')
add_table(
    ['입력 항목', '예시', '활용 단계'],
    [
        ['화물 종류', '일반화물 / 냉장화물 / 위험물 / 2차전지 등 8종', 'Step 3 창고 유형 필터'],
        ['CBM (용적)', '15.0 CBM', '비용 계산 기준'],
        ['항로', '부산→로테르담 (5개 선택지)', '시나리오 영향 범위 판단'],
        ['납기일', '2026-06-01', '지연 여유 시간 산출'],
        ['집화일', '2026-05-20', 'Phase 1 루티 시작점'],
        ['출발 권역', '경기남부 / 부산 등', 'ODCY·CY 거리 계산'],
    ],
    [3.5, 6, 5],
)

# Step 2
h2('Step 2 — MRI 분석')
box('Maritime Risk Index: 실시간 뉴스 + 운임지수 + 물동량을 5차원 AHP로 종합한 0~1 리스크 지수입니다.', 'FFE0B2')

h3('MRI 산출 공식 (AHP 일관성 CR=3.1%)')
box('MRI  =  0.431 x G  +  0.182 x D  +  0.253 x F  +  0.090 x V  +  0.044 x P', 'FFF9C4')

add_table(
    ['차원', '이름', '계산 방법', '포화점', '가중치'],
    [
        ['G', '지정학·항로', '뉴스 내 지정학 비중 / 0.25', '25%', '43.1%'],
        ['D', '지연·운항', 'G x 1.0 + 기상 x 0.36 (프록시)', '-', '18.2%'],
        ['F', '운임 변동', '뉴스비중/0.20 x 50% + KCCI변동/0.15 x 50%', '15~20%', '25.3%'],
        ['V', '통행량', 'KCCI 물동량 감소율 / 0.10', '10% 감소', '9.0%'],
        ['P', '항만·통상', '뉴스 내 항만·관세 비중 / 0.20', '20%', '4.4%'],
    ],
    [1.5, 3, 5.5, 2.5, 2],
)

note('포화점 설계 이유: RSS 전체 기사 중 단일 카테고리가 25% 이상 차지하기 어렵습니다. 포화점 없이 비율을 그대로 쓰면 시뮬레이션(G=0.5~0.7)과 5배 괴리가 발생합니다.')

h3('MRI 등급 및 시나리오 분류')
add_table(
    ['등급', 'MRI 범위', '시나리오 ID', '대응 정책', '지연', '운임 변동'],
    [
        ['정상 (초록)', '< 0.3', 'A_NORMAL', 'AS_PLANNED (기존 계획 유지)', '-', '-'],
        ['주의 (노랑)', '0.3 ~ 0.5', 'D_DELAY', 'SHIFT_PICKUP (집화 일정 조정)', '+3일', '+2%'],
        ['경계 (주황)', '>= 0.5 + 기상', 'C_WEATHER', 'HOLDBACK / 콜드체인 우선', '+5일', '+5%'],
        ['위험 (빨강)', '>= 0.7 + 지정학', 'B_GEOPOLITICAL', 'REROUTE + 케이프타운 우회', '+14일', '+30%'],
        ['특수', 'cancel_flag', 'E_CANCELLATION', 'REGROUP_REMAINING (재구성)', '-', '-'],
    ],
    [2, 2.5, 3.5, 5, 1.5, 2.5],
)

h3('Step 2 추가 정보')
bullet('과거 유사사례 매칭: 에버기븐 좌초, 홍해 후티 공격, 태풍 힌남노 등 7개 실제 사건 DB → 평균 지연일·운임 변동 제시')
bullet('LSTM 물동량 예측: 부산항 2020~2025 실데이터 학습 → 3개월 TEU 예측 (MAPE 9.4%)')
bullet('Gemini LLM 자동 보고서: MRI·시나리오·유사사례를 자연어 보고서로 자동 생성')

doc.add_page_break()

# Step 3
h2('Step 3 — 창고·ODCY 추천 (MRI >= 0.5 시 활성화)')
box('MRI 0.5 미만이면 이 단계는 건너뜁니다. 화주의 최종 선택을 돕는 추천 단계이며, 플랫폼은 강제하지 않습니다.', 'D4EDDA')

h3('창고 탐색 우선순위 (3단계)')
add_table(
    ['우선순위', '소스', '내용', '창고 수'],
    [
        ['1순위', 'NLIC 국가물류통합정보센터', '정부 공인 부산 물류창고 DB (지오코딩 완료)', '439개'],
        ['2순위', '카카오 Local API', '항만 반경 15km 실시간 키워드 검색', '실시간'],
        ['3순위', '내장 시뮬 DB', 'NLIC 없을 때 폴백 (9개 대표 창고)', '9개'],
    ],
    [2, 5, 6, 2],
)

h3('창고 유형별 화물 매핑')
add_table(
    ['화물 종류', '필수 키워드', '냉동 필요', '위험물 허가'],
    [
        ['일반화물', '물류창고, 물류센터, CY, 보세창고', '불필요', '불필요'],
        ['냉장화물', '냉동창고, 냉장창고, 저온물류', '0~10°C', '불필요'],
        ['냉동화물', '냉동창고, 저온물류', '-25~-18°C', '불필요'],
        ['위험물', '보세창고, 위험물창고', '불필요', '필수'],
        ['2차전지', '보세창고, 위험물창고', '15~25°C', '필수 (IMDG Class 9)'],
        ['전자제품', '보세창고, 물류센터', '불필요 (정온)', '불필요'],
    ],
    [3, 5, 3, 3],
)

h3('A/B/C/D 4가지 옵션 비교')
add_table(
    ['옵션', '전략', '비용 특징', '추천 상황'],
    [
        ['A — 직송·리스크 감수', '창고 없이 바로 CY 직송', '가장 저렴, 리스크 있음', 'MRI < 0.5 경계선'],
        ['B — 최단거리 창고', '항만에서 거리 최단 창고', '운송비 최소', '긴급 보관 필요 시'],
        ['C — 최저비용 창고', '보관료+운송비 합산 최소', '총비용 최적화', '일정 여유 있을 때'],
        ['D — 종합 권장', '거리+시간+시설 종합 점수 1위', '최적 균형', '불확실성 높을 때 권장'],
    ],
    [3, 4.5, 4, 4],
)

note('플랫폼은 D 옵션을 기본 추천하나 최종 결정은 화주에게 있습니다. 루티 JSON은 화주가 선택한 옵션 기준으로 생성됩니다.')

# Step 4
h2('Step 4 — 루티 JSON 생성')
box('화주가 선택한 창고와 경로를 위밋 루티(ROOUTY)/루티프로 API 입력 형식 JSON으로 자동 변환합니다.', 'EDE7F6')

add_table(
    ['Phase', '구간', '내용', '파일명 패턴'],
    [
        ['Phase 1', '출발지 → 창고·ODCY', '화주 공장/창고에서 항만 인근 보관소까지 배차', 'EG-YYYYMMDD-STORAGE-SH-xxx.json'],
        ['Phase 2', '창고·ODCY → CY', '선적 재개 시 보관소에서 컨테이너야드까지 배차', 'EG-YYYYMMDD-PHASE2-SH-xxx.json'],
    ],
    [2, 5, 7, 5],
)

note('루티 연동 상태: 현재 simulation_mode. 루티 API 발급 시 integration_status를 live_api로 변경하면 즉시 연동됩니다.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. 데이터 파이프라인
# ════════════════════════════════════════════════════════════
h1('3. 데이터 파이프라인 및 소스')

arrow_row([
    ('실시간\n뉴스 RSS', 'FFE0CC'),
    ('NLP\n키워드 분류', 'FFF3E0'),
    ('MRI\nAHP 산출', 'FFF9C4'),
    ('시나리오\n분류', 'D4EDDA'),
    ('루티 JSON\n출력', 'E8EAF6'),
])

add_table(
    ['데이터', '소스', '수집 주기', '없을 때 대체'],
    [
        ['해사 뉴스', 'gCaptain RSS, Splash247 RSS, 한국해운신문 RSS', '실시간 (30일치, 소스당 30건)', '시뮬 뉴스 3건 사용'],
        ['KCCI 운임지수', '한국해양진흥공사 XLS 파일', '주간 (수동 업데이트)', '기본값 0.20 사용'],
        ['부산항 물동량', 'BPA API (2020~2024) + 홈페이지 Excel (2025)', '월간', 'BPA 직접 다운로드'],
        ['유가 (브렌트유)', 'Yahoo Finance BZ=F → FRED → EIA 순', '일간 자동', 'Yahoo Finance 자동 대체'],
        ['환율 (원/달러)', 'ECOS(한국은행) → frankfurter.app', '일간 자동', 'frankfurter.app 자동 대체'],
        ['창고·ODCY', 'NLIC DB(439개) + 카카오 Local API', '정적(NLIC) + 실시간(카카오)', '내장 시뮬 DB 9개'],
        ['경로·거리', '카카오모빌리티 길찾기 API', '요청 시 실시간', 'Haversine 직선거리 추정'],
        ['LLM 보고서', 'Google Gemini 1.5 Flash (무료, 1500회/일)', '요청 시', 'Claude Haiku (유료) 대체'],
    ],
    [3, 6, 3.5, 4],
)

h2('뉴스 NLP 분류 체계')
add_table(
    ['카테고리', '한글 키워드 (예시)', '영어 키워드 (예시)', 'MRI 반영 차원'],
    [
        ['지정학분쟁', '후티, 홍해, 수에즈, 이란', 'houthi, suez, hormuz, iran', 'G (지정학)'],
        ['기상재해', '태풍, 폭풍, 결빙, 항로폐쇄', 'typhoon, storm, ice, canal drought', 'D (지연)'],
        ['운임변동', '운임급등, SCFI, 할증료', 'freight rate, SCFI, surcharge', 'F (운임)'],
        ['항만파업', '파업, 폐쇄, 혼잡, 적체', 'strike, closure, congestion', 'P (항만)'],
        ['관세정책', '관세, 무역제재, 트럼프', 'tariff, trump, trade war', 'P (항만)'],
        ['공급망이슈', '지연, 부족, 용량', 'delay, shortage, capacity', 'V (통행량)'],
    ],
    [3, 5, 5, 3],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. 시스템 아키텍처
# ════════════════════════════════════════════════════════════
h1('4. 시스템 아키텍처')

h2('4-1. 구성 요소')
add_table(
    ['레이어', '구성 요소', '기술 스택', '역할'],
    [
        ['프론트엔드', 'Lovable (React)', 'React + Vite + TypeScript', '화주 UI, 4단계 화면 구현'],
        ['백엔드 API', 'api.py (FastAPI)', 'Python + FastAPI + Uvicorn', 'REST API 8개 엔드포인트'],
        ['분석 엔진', 'src/ (15개 모듈)', 'Python (numpy, pandas, torch)', 'MRI, 시나리오, 창고, 루티'],
        ['데이터', 'data/ 폴더', 'JSON, XLS, CSV, PNG', '창고DB, 캐시, 시각화'],
        ['클라우드', 'Render (싱가포르)', 'Docker + Python 3.10+', 'FastAPI 서버 배포'],
        ['노트북', 'wemeet_v4_main.ipynb', 'Jupyter (29셀)', '발표용 시연'],
    ],
    [3, 3.5, 5, 5],
)

h2('4-2. API 엔드포인트')
add_table(
    ['Method', 'Endpoint', '기능', 'Step'],
    [
        ['GET',  '/api/health',              '서버 상태 확인',                          '-'],
        ['GET',  '/api/mri',                 '현재 MRI + 등급 + 시나리오',              'Step 2'],
        ['GET',  '/api/mri/similar-events',  '과거 유사사례 3건 + 평균 지연·운임',      'Step 2'],
        ['GET',  '/api/mri/lstm-forecast',   'LSTM 3개월 물동량 예측',                  'Step 2'],
        ['GET',  '/api/routes',              '이용 가능 항로 5개 목록',                 'Step 1'],
        ['POST', '/api/shipment/register',   '화주 출하 등록 + 영향 분석',              'Step 1~2'],
        ['POST', '/api/warehouse/recommend', '창고·ODCY 추천 + 4옵션 비용 비교',        'Step 3'],
        ['POST', '/api/routy/generate',      'Phase1+2 루티 JSON 자동 생성',            'Step 4'],
    ],
    [1.5, 5, 6, 2],
)

h2('4-3. 배포 구조')
box('로컬 (발표용): uvicorn api:app --reload --port 8000  |  Streamlit: streamlit run app.py', 'F5F5F5')
box('서버 배포: Render (Singapore) — render.yaml 자동 설정  |  URL: https://wemeet-api-dchk.onrender.com', 'D6E4F0')
box('프론트엔드: Lovable 프로젝트  |  VITE_API_BASE_URL = Render URL 설정', 'D4EDDA')
note('Render 무료 플랜: 15분 미사용 시 슬립. 발표 전 /api/health 접속으로 서버를 미리 깨워두세요.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. 파일 구조
# ════════════════════════════════════════════════════════════
h1('5. 주요 파일 구조 및 역할')

add_table(
    ['파일 / 폴더', '역할', '핵심 함수·내용'],
    [
        ['api.py', 'FastAPI 백엔드 (Lovable 연동)', '8개 REST API 엔드포인트'],
        ['app.py', 'Streamlit 웹앱 (4탭)', '화주입력 / MRI / 창고 / 루티 탭'],
        ['src/config.py', '시나리오·AHP·키워드 상수', 'SCENARIOS, MRI_AHP_WEIGHTS, ROUTE_INFO'],
        ['src/mri_engine.py', 'MRI 5차원 AHP 산출', 'calc_today_mri(), build_mri_series()'],
        ['src/nlp_classifier.py', '뉴스 키워드 분류 (한글+영어)', 'classify_news_df(), top_category()'],
        ['src/scenario_engine.py', '시나리오 자동 분류 + 영향 분석', 'auto_classify_scenario(), analyze_impact()'],
        ['src/historical_matcher.py', '과거 유사사례 매칭 (7개 사건)', 'find_similar_events()'],
        ['src/lstm_forecaster.py', 'LSTM 부산항 물동량 예측', 'train_and_forecast(), build_main_df()'],
        ['src/odcy_recommender.py', 'NLIC+카카오 API 창고 탐색', 'recommend_storage(), _load_nlic_db()'],
        ['src/option_presenter.py', 'A/B/C/D 4가지 옵션 비용 산출', 'generate_four_options()'],
        ['src/storage_routy_adapter.py', 'Phase 1/2 루티 JSON 생성', 'generate_storage_routy_json()'],
        ['src/real_data_fetcher.py', 'RSS 뉴스·유가·환율 실시간 수집', 'fetch_maritime_news() — 30일치'],
        ['src/llm_reporter.py', 'Gemini/Claude 자동 보고서 생성', 'generate_risk_report()'],
        ['data/nlic_warehouses.json', 'NLIC 부산 물류창고 DB', '439개 창고 (좌표·면적·화물유형)'],
        ['data/lstm_cache.json', 'LSTM 사전계산 캐시', 'Render 서버 배포용 (torch 없이 작동)'],
        ['notebooks/wemeet_v4_main.ipynb', '발표용 메인 노트북', '29셀 순차 실행 — 전 기능 시연'],
    ],
    [5, 5, 6],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. 시나리오별 상세 대응 흐름
# ════════════════════════════════════════════════════════════
h1('6. 시나리오별 상세 대응 흐름')

h2('시나리오 B — 지정학 분쟁 (홍해 사태 기준)')
add_table(
    ['항목', '내용'],
    [
        ['트리거 조건', 'MRI >= 0.7 + 뉴스 키워드: houthi / suez / hormuz / iran / tariff'],
        ['대응 정책', 'REROUTE_AND_HOLDBACK — 케이프타운 우회 + 출하 보류'],
        ['지연 일수', '+14일 (케이프타운 우회 실제 소요일 — 변경 금지)'],
        ['운임 변동', '+30% (홍해 사태 실제 운임 증가율 — 변경 금지)'],
        ['참고 근거', 'UNCTAD 2024: 수에즈 통항 42~90% 감소 / Drewry WCI 2024'],
        ['우회 경로', '수에즈 운하 대신 케이프타운 → +11,000 km, +14일'],
        ['세부 시나리오', 'B1(홍해/수에즈), B2(호르무즈), B3(미중 관세)로 자동 세분화'],
    ],
    [4, 12],
)

h2('시나리오 C — 기상 악화 (태풍 기준)')
add_table(
    ['항목', '내용'],
    [
        ['트리거 조건', 'MRI >= 0.5 + 뉴스 키워드: typhoon / storm / weather / canal drought'],
        ['대응 정책', 'HOLDBACK_NORMAL_RUSH_COLD — 보류 + 콜드체인 화물 우선 처리'],
        ['지연 일수', '+5일 (기상 대기 시간)'],
        ['운임 변동', '+5% (긴급 배차 비용)'],
        ['참고 근거', '태풍 힌남노(2022): 부산항 5일 폐쇄, 운임 +6%'],
    ],
    [4, 12],
)

h2('시나리오 D — 일반 지연')
add_table(
    ['항목', '내용'],
    [
        ['트리거 조건', 'MRI 0.3 ~ 0.5'],
        ['대응 정책', 'SHIFT_PICKUP — 집화 일정 3일 조정'],
        ['지연 일수', '+3일'],
        ['운임 변동', '+2% (소폭 조정)'],
    ],
    [4, 12],
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. 발표 핵심 메시지
# ════════════════════════════════════════════════════════════
h1('7. 발표 핵심 메시지 및 차별점')

h2('7-1. 핵심 차별점 3가지')
box('[1] 선제적 리스크 의사결정: MRI 5차원 AHP로 리스크 발생 전 대응 옵션 자동 제시 — 기존 플랫폼은 사후 알림만 제공', 'D6E4F0')
box('[2] 정부 공인 창고 데이터: NLIC 국가물류통합정보센터 439개 창고 DB + 카카오 실시간 검색 병행 — 단순 구글 지도 검색과 차별화', 'D4EDDA')
box('[3] 루티 직접 연동: 화주의 결정이 즉시 루티 배차 JSON으로 변환 — 의사결정과 실행의 원스톱 연결', 'EDE7F6')

h2('7-2. 트레드링스와의 관계')
note('트레드링스·Flexport는 경쟁 상대가 아닌 보완재입니다. 해상 구간 가시성은 기존 플랫폼을 활용하고, 본 플랫폼은 선제적 국내 운송 재조정과 의사결정 지원에 집중합니다.')

h2('7-3. 실데이터 기반 신뢰성')
add_table(
    ['데이터', '출처', '검증 지표'],
    [
        ['부산항 물동량 LSTM', 'BPA(부산항만공사) 2020~2025 실데이터', 'MAPE 9.4% (시간순 분할)'],
        ['창고 DB', 'NLIC 국가물류통합정보센터 공공데이터', '439개 실제 등록 창고'],
        ['시나리오 파라미터', 'UNCTAD 2024, Drewry WCI, BPA 운영보고서', '+14일, +30% 근거 있음'],
        ['운임 데이터', 'KCCI(한국해양진흥공사) XLS 원본', '주간 업데이트'],
        ['AHP 가중치', '쌍대비교 행렬 CR=3.1% < 10%', '통계적 일관성 검증 완료'],
    ],
    [4, 6, 5],
)

# ════════════════════════════════════════════════════════════
# 저장
# ════════════════════════════════════════════════════════════
doc.save(OUT)
print(f'[OK] 저장 완료: {OUT}')
