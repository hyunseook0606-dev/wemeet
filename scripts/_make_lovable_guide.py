# -*- coding: utf-8 -*-
"""Lovable 연동 가이드 DOCX 생성 (상세 버전)."""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print('python-docx 설치 필요: pip install python-docx')
    sys.exit(1)

OUT = Path(__file__).parent.parent / 'Lovable_연동_개발가이드.docx'
doc = Document()

# 페이지 여백
s = doc.sections[0]
s.left_margin = s.right_margin = Cm(2.5)
s.top_margin  = s.bottom_margin = Cm(2.0)

# ── 스타일 헬퍼 ────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(10)

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'E8F4FD')
    p._p.get_or_add_pPr().append(shading)
    r = p.add_run('  ' + text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFF3CD')
    p._p.get_or_add_pPr().append(shading)
    r = p.add_run('  ⚠ ' + text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x85, 0x65, 0x04)

def bullet(text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text).font.size = Pt(10)

def numberd(text, num, indent=0.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{num}.  ')
    r1.bold = True; r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.add_run(text).font.size = Pt(10)

def code(text, bg='F2F2F2'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), bg)
    p._p.get_or_add_pPr().append(shading)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9)
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E75B6')
        c._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        fill = 'F0F7FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = val
            c.paragraphs[0].runs[0].font.size = Pt(9)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            c._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    p.add_run('─' * 75).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ════════════════════════════════════════════════════════════════════════════
# 표지
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('위밋모빌리티 해상 리스크 플랫폼')
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Lovable 프론트엔드 연동 개발 가이드')
r.bold = True; r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('작성일: 2026년 5월  |  버전: 2.0  |  대상: 프론트엔드 개발 팀원').font.size = Pt(10)

doc.add_paragraph()
tip('이 문서는 Lovable, React, API 경험이 전혀 없는 분도 처음부터 따라할 수 있도록 작성되었습니다.')
doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 목차
# ════════════════════════════════════════════════════════════════════════════
h1('목차')
items = [
    ('1', '전체 구조 이해하기'),
    ('2', '백엔드 서버 실행하기 (개발팀 도움 필요)'),
    ('3', 'Lovable 시작하기'),
    ('4', '화면 1 — 대시보드 만들기'),
    ('5', '화면 2 — 출하 등록 폼 만들기'),
    ('6', '화면 3 — 창고 추천 + 옵션 비교 만들기'),
    ('7', 'API 연결 코드 작성법'),
    ('8', '전체 테스트 방법'),
    ('9', '배포하기 (Railway + Lovable)'),
    ('10', '트러블슈팅'),
]
for num, title in items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{num}.  ')
    r1.bold = True; r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.add_run(title).font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. 전체 구조 이해
# ════════════════════════════════════════════════════════════════════════════
h1('1. 전체 구조 이해하기')
body('개발을 시작하기 전에 전체 그림을 이해하면 훨씬 쉽습니다.')
doc.add_paragraph()

h2('1-1. 우리가 만드는 것')
body('위밋모빌리티 플랫폼은 크게 두 부분으로 나뉩니다:')
bullet('백엔드 (Python): 이미 완성되어 있습니다. 해상 리스크를 계산하고, 창고를 탐색하고, 비용을 분석합니다.')
bullet('프론트엔드 (Lovable): 여러분이 만들 부분입니다. 화주(고객)가 보는 화면입니다.')
doc.add_paragraph()

h2('1-2. 두 부분이 대화하는 방법')
body('프론트엔드와 백엔드는 "API"라는 방식으로 대화합니다. 쉽게 말하면 이렇습니다:')
code(
    '예시:\n'
    '  화주가 "출하 등록" 버튼을 클릭\n'
    '      ↓\n'
    '  Lovable 화면이 백엔드에 요청을 보냄:\n'
    '  "경기식품, 냉장화물 15CBM, 부산→로테르담으로 보내고 싶어"\n'
    '      ↓\n'
    '  백엔드가 계산 후 응답:\n'
    '  "예상 운임 $675, 14일 지연 예상, 창고 보관 옵션 있음"\n'
    '      ↓\n'
    '  Lovable 화면에 결과 표시'
)

h2('1-3. 아키텍처 그림')
code(
    '┌─────────────────────────────────┐\n'
    '│   Lovable 프론트엔드            │  ← 여러분이 만드는 부분\n'
    '│   React + TypeScript + Tailwind │\n'
    '│                                 │\n'
    '│   화면 1: 대시보드 (/dashboard) │\n'
    '│   화면 2: 출하등록 (/shipment)  │\n'
    '│   화면 3: 창고추천 (/warehouse) │\n'
    '└────────────┬────────────────────┘\n'
    '             │  HTTP 요청/응답 (JSON)\n'
    '┌────────────▼────────────────────┐\n'
    '│   FastAPI 백엔드 (api.py)       │  ← 이미 완성\n'
    '│   http://localhost:8000         │\n'
    '│                                 │\n'
    '│   /api/mri           MRI 계산   │\n'
    '│   /api/shipment      출하 분석  │\n'
    '│   /api/warehouse     창고 탐색  │\n'
    '│   /api/routy         JSON 생성  │\n'
    '└────────────┬────────────────────┘\n'
    '             │\n'
    '┌────────────▼────────────────────┐\n'
    '│   Python 모듈 (src/)            │  ← 이미 완성\n'
    '│   MRI 엔진, LSTM, 카카오 API    │\n'
    '└─────────────────────────────────┘'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. 백엔드 실행
# ════════════════════════════════════════════════════════════════════════════
h1('2. 백엔드 서버 실행하기')
warn('이 단계는 Python 백엔드 개발자(또는 팀원)가 수행합니다. 프론트엔드 개발자는 백엔드가 실행된 상태에서 시작하면 됩니다.')

h2('2-1. 백엔드 실행 방법')
body('Python이 설치된 PC에서 아래 명령어를 실행합니다:')
code(
    '# 1단계: 패키지 설치 (최초 1회만)\n'
    'pip install fastapi uvicorn python-dotenv\n\n'
    '# 2단계: 프로젝트 폴더로 이동\n'
    'cd "C:\\Users\\USER\\Desktop\\vs code\\vs code"\n\n'
    '# 3단계: 서버 시작\n'
    'uvicorn api:app --reload --port 8000'
)

body('아래와 같은 메시지가 나오면 성공입니다:')
code(
    'INFO:     Uvicorn running on http://127.0.0.1:8000 (Press CTRL+C to quit)\n'
    'INFO:     Started reloader process\n'
    'INFO:     Application startup complete.'
)

h2('2-2. API 동작 확인')
body('백엔드가 실행 중일 때 브라우저에서 아래 주소를 열어보세요:')
bullet('http://localhost:8000/docs → Swagger UI (API 문서 + 테스트 화면)')
bullet('http://localhost:8000/api/health → {"status": "ok"} 응답이 나오면 정상')
bullet('http://localhost:8000/api/mri → MRI 점수가 JSON으로 나오면 정상')

tip('/docs 페이지에서 각 API를 클릭 → "Try it out" → "Execute"로 직접 테스트할 수 있습니다.')

h2('2-3. 프론트엔드 개발자에게 알려줄 정보')
body('백엔드가 실행되면 프론트엔드 개발자에게 아래 주소를 알려주세요:')
code('API 주소: http://localhost:8000')
body('같은 네트워크(같은 WiFi)에 있다면 localhost 대신 PC의 IP 주소를 사용합니다:')
code(
    '# PC의 IP 주소 확인 방법 (Windows PowerShell)\n'
    'ipconfig\n'
    '# IPv4 Address 항목 확인 (예: 192.168.1.100)\n\n'
    '# 그러면 API 주소는:\n'
    'http://192.168.1.100:8000'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. Lovable 시작하기
# ════════════════════════════════════════════════════════════════════════════
h1('3. Lovable 시작하기')

h2('3-1. Lovable이란?')
body('Lovable(lovable.dev)은 AI에게 말로 설명하면 자동으로 React 웹앱을 만들어주는 서비스입니다. '
     '코딩 지식이 없어도 사용할 수 있으며, 만들어진 코드는 GitHub에 저장되고 자동으로 배포됩니다.')

h2('3-2. 계정 만들기')
numberd('https://lovable.dev 접속', 1)
numberd('"Sign up" 클릭 → GitHub 계정으로 가입 (권장)', 2)
numberd('이메일 인증 완료', 3)
tip('GitHub 계정이 없다면 github.com에서 먼저 계정을 만드세요 (무료).')

h2('3-3. 새 프로젝트 만들기')
numberd('"New Project" 클릭', 1)
numberd('프로젝트 이름 입력: wemet-platform', 2)
numberd('"Create Project" 클릭', 3)
numberd('채팅창이 열리면 준비 완료', 4)

h2('3-4. 첫 번째 프롬프트 (전체 구조 생성)')
body('채팅창에 아래 내용을 그대로 복붙하고 전송하세요.')
tip('한 번에 모든 화면을 만들려고 하지 말고, 아래 프롬프트로 기본 구조를 먼저 잡으세요.')

code(
    'Create a Korean maritime risk platform called "위밋모빌리티".\n'
    'Tech stack: React + TypeScript + Tailwind CSS.\n\n'
    'Create 3 pages with navigation:\n'
    '1. /dashboard - MRI Dashboard (해상 리스크 모니터링)\n'
    '2. /shipment - Shipment Form (출하 등록)\n'
    '3. /warehouse - Warehouse Recommendation (창고 추천)\n\n'
    'Header: Navy blue (#1F4E79) with company name "위밋모빌리티"\n'
    'Navigation menu with 3 links\n'
    'Each page should have a title and placeholder content for now\n\n'
    'API_BASE_URL constant: "http://localhost:8000"\n'
    'Create a file src/config/api.ts with:\n'
    'export const API_BASE = "http://localhost:8000";'
)

body('Lovable이 기본 구조를 만들어주면 "Preview" 버튼으로 결과를 확인하세요.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 4. 화면 1 — 대시보드
# ════════════════════════════════════════════════════════════════════════════
h1('4. 화면 1 — 대시보드 만들기 (/dashboard)')
body('대시보드는 현재 해상 리스크 현황을 한눈에 보여주는 화면입니다.')

h2('4-1. 화면 구성')
table(
    ['섹션', '내용', '데이터 출처'],
    [
        ['MRI 점수 카드', 'MRI 수치 + 색상 배지 (위험/경계/주의/정상)', 'GET /api/mri'],
        ['MRI 설명', '"MRI(해상 리스크 지수)란?" 접기/펼치기', '정적 텍스트'],
        ['과거 유사사례', '평균 지연일 + 평균 운임 상승률', 'GET /api/mri/similar-events'],
        ['LSTM 예측', '향후 3개월 부산항 물동량 바차트', 'GET /api/mri/lstm-forecast'],
        ['실시간 뉴스', '최신 해사 뉴스 목록 (선택)', '(선택)'],
    ],
    col_widths=[3.5, 5, 5]
)

h2('4-2. Lovable 프롬프트')
body('Lovable 채팅창에 아래 내용을 입력하세요:')
code(
    'Update the /dashboard page with these sections:\n\n'
    '1. TOP SECTION - MRI Score Card\n'
    'Fetch GET http://localhost:8000/api/mri on page load\n'
    'Show:\n'
    '- Large number: mri value (e.g. "0.717")\n'
    '- Badge with grade text (e.g. "경계")\n'
    '- Badge color based on grade:\n'
    '  "위험" -> red (#EF5350)\n'
    '  "경계" -> orange (#FF7043)\n'
    '  "주의" -> yellow (#FFA726)\n'
    '  "정상" -> green (#66BB6A)\n'
    '- Text: "주요 리스크: {category}"\n'
    '- Text: "자동 분류 시나리오: {scenario_name}"\n'
    '- Show loading spinner while fetching\n\n'
    '2. ACCORDION - MRI란? (Expandable section)\n'
    'Title: "MRI(해상 리스크 지수)란?"\n'
    'Default: expanded if MRI >= 0.5, collapsed otherwise\n'
    'Content: Table with 4 rows:\n'
    '  위험 | 0.8 이상 | 항로 봉쇄·전쟁 수준\n'
    '  경계 | 0.6~0.8  | 심각한 운임·지연 예상\n'
    '  주의 | 0.3~0.6  | 파업·혼잡·소규모 분쟁\n'
    '  정상 | 0.3 미만 | 정상 운항\n\n'
    '3. SIMILAR EVENTS SECTION\n'
    'Title: "과거 유사사례 분석"\n'
    'Fetch GET http://localhost:8000/api/mri/similar-events\n'
    'Show 2 metric cards:\n'
    '  - "유사사례 평균 지연" : avg_delay + "일"\n'
    '  - "유사사례 평균 운임 상승" : "+" + avg_freight + "%"\n'
    'List of events below (rank, name, year)\n\n'
    '4. LSTM FORECAST\n'
    'Title: "부산항 물동량 예측 (LSTM)"\n'
    'Fetch GET http://localhost:8000/api/mri/lstm-forecast\n'
    'Show 3 metric cards with month and teu_10k value\n'
    'Add caption: "단위: 만 TEU/월 | MAPE: {mape}%"'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. 화면 2 — 출하 등록
# ════════════════════════════════════════════════════════════════════════════
h1('5. 화면 2 — 출하 등록 폼 만들기 (/shipment)')
body('화주가 화물 정보를 입력하고 리스크 영향을 즉시 확인하는 화면입니다.')

h2('5-1. 입력 필드')
table(
    ['필드', '타입', 'API 파라미터명', '예시값'],
    [
        ['화주명 (회사명)', '텍스트 입력', 'company', '(주)경기식품'],
        ['화물 유형', '드롭다운', 'cargo_type', '일반화물 / 냉장화물 / 위험물'],
        ['화물 용량', '숫자 입력', 'cbm', '15'],
        ['희망 항로', '드롭다운 (API에서 로드)', 'route', '부산→로테르담'],
        ['집화 권역', '드롭다운', 'region', '경기남부'],
        ['희망 집화일', '날짜 선택', 'pickup_date', '2026-05-20'],
        ['납기 여유', '드롭다운', 'deadline_days', '7 / 10 / 14 / 21'],
        ['긴급 화물', '체크박스', 'urgent', 'false'],
    ],
    col_widths=[3.5, 2.5, 3.5, 4]
)

h2('5-2. 제출 후 표시할 결과')
table(
    ['항목', '표시 내용', '조건'],
    [
        ['예상 운임', '$675 (USD)', '항상'],
        ['지연 예상', '+14일', '항상'],
        ['운임 변화', '+$203', '항상'],
        ['납기 위반 여부', '⚠ 위험 / ✅ 안전', '항상'],
        ['MRI 맥락', '현재 MRI + 과거 유사사례', '항상'],
        ['창고 보관 옵션 버튼', '"창고 보관 옵션 보기" 버튼', 'show_warehouse === true 일 때만'],
    ],
    col_widths=[3.5, 4, 5]
)

h2('5-3. Lovable 프롬프트')
code(
    'Update the /shipment page:\n\n'
    'PART 1 - Form\n'
    'Create a form with these fields:\n'
    '- company: text input (placeholder: "(주)경기식품")\n'
    '- cargo_type: select with options ["일반화물", "냉장화물", "위험물"]\n'
    '- cbm: number input (min:1, max:100, default:15)\n'
    '- route: select - fetch options from GET http://localhost:8000/api/routes\n'
    '  use routes[].name for display, routes[].name for value\n'
    '- region: select ["경기남부","경기북부","충청","경상남부","경상북부"]\n'
    '- pickup_date: date input\n'
    '- deadline_days: select [7, 10, 14, 21]\n'
    '- urgent: checkbox\n'
    'Submit button: "출하 등록 & 분석" (primary, full width)\n\n'
    'PART 2 - Result (show after submit)\n'
    'POST to http://localhost:8000/api/shipment/register\n'
    'Request body: { company, cargo_type, cbm, route, pickup_date,\n'
    '                deadline_days, region, urgent }\n'
    'pickup_date format: "YYYY-MM-DD"\n\n'
    'Show result card with:\n'
    '- 4 metric cards: estimated_cost ($), delay_days (+Xd),\n'
    '  cost_delta ($+/-), deadline_violated (warning/ok)\n'
    '- MRI context box (blue background):\n'
    '  "현재 MRI {mri} — {scenario_name}"\n'
    '- If response.requires_priority === true:\n'
    '  show info box "⭐ 우선처리 대상"\n'
    '- If response.requires_holdback === true:\n'
    '  show warning box "◐ 항만 반입 보류 권고"\n'
    '- If response.show_warehouse === true:\n'
    '  show primary button "🏭 창고 보관 옵션 보기"\n'
    '  On click: navigate to /warehouse\n'
    '  Pass as query params: port_name, cargo_type, cbm,\n'
    '  delay_days, freight_usd=estimated_cost'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. 화면 3 — 창고 추천
# ════════════════════════════════════════════════════════════════════════════
h1('6. 화면 3 — 창고 추천 + 옵션 비교 (/warehouse)')
body('MRI가 높을 때(≥0.5) 항만 주변 창고를 자동으로 찾고, 4가지 대응 옵션의 비용을 비교하는 화면입니다.')
warn('이 화면은 화주에게 "추천"만 제공합니다. 최종 결정은 화주가 직접 합니다.')

h2('6-1. 화면 구성')
table(
    ['섹션', '내용', '데이터'],
    [
        ['헤더', '현재 MRI + "추천이며 최종결정은 화주에게" 안내', '쿼리 파라미터'],
        ['창고 목록', '3개 창고 카드 (이름/주소/거리/운영시간)', 'POST /api/warehouse/recommend'],
        ['옵션 비교', 'A/B/C/D 4가지 비용 카드', '위와 동일'],
        ['D안 상세', '권장 옵션 세부 비용 내역', '위와 동일'],
        ['JSON 생성', '"루티 JSON 생성" 버튼 → 파일 다운로드', 'POST /api/routy/generate'],
    ],
    col_widths=[3, 6, 4.5]
)

h2('6-2. Lovable 프롬프트')
code(
    'Update /warehouse page:\n\n'
    'Read query params from URL:\n'
    'port_name, cargo_type, cbm, delay_days, freight_usd\n\n'
    'SECTION 1 - Header\n'
    'Show: "Step 3. 창고·ODCY 자동 탐색"\n'
    'Show caption: "MRI 기반 창고 보관 옵션 — 최종 결정은 화주님께 있습니다."\n\n'
    'SECTION 2 - Fetch warehouses\n'
    'POST http://localhost:8000/api/warehouse/recommend\n'
    'Body: { port_name, cargo_type, cbm: parseFloat(cbm),\n'
    '        mri_score: 0.7, delay_days: parseInt(delay_days),\n'
    '        freight_usd: parseInt(freight_usd) }\n'
    'Show loading spinner while fetching\n\n'
    'SECTION 3 - Warehouse Cards\n'
    'For each item in response.warehouses (max 3):\n'
    '  Card with:\n'
    '  - Title: warehouse.name\n'
    '  - Address: warehouse.address\n'
    '  - "거리: {distance_km}km / {duration_min}분"\n'
    '  - "운영: {operating_hours}"\n'
    '  - If warehouse.special_notes: yellow warning box\n\n'
    'SECTION 4 - Option Comparison\n'
    'Title: "A/B/C/D 4가지 대응 옵션"\n'
    'Subtitle: "모든 비용은 USD 기준"\n'
    'For each option in response.options:\n'
    '  Card (4 in a row) with:\n'
    '  - Badge: option.id + "안" (A안/B안/C안/D안)\n'
    '  - Star badge on D: "★ 권장"\n'
    '  - Total: "$" + option.total_usd\n'
    '  - If i>0: savings in green "-$" + option.savings_usd\n'
    '  - option.name (smaller text)\n'
    '  D안 card should have blue border highlight\n\n'
    'SECTION 5 - D Option Detail (collapsible)\n'
    'Title: "★ D안 (권장) 상세 비용"\n'
    'Show breakdown for D option:\n'
    '  해상 운임: $X\n'
    '  루티 운송 P1: $X\n'
    '  루티 운송 P2: $X\n'
    '  창고 대여비: $X\n'
    '  창고 계약비: $X\n'
    '  합계: $X\n'
    '  A안 대비 절약: $X\n\n'
    'SECTION 6 - JSON Generate Button\n'
    'Button: "📤 루티 연계 JSON 생성 (D안 선택)"\n'
    'On click: POST http://localhost:8000/api/routy/generate\n'
    'Body (use D option warehouse and query params):\n'
    '{\n'
    '  shipment_id: "SH-" + Date.now(),\n'
    '  company: "화주",\n'
    '  region: "경기남부",\n'
    '  cargo_type: cargo_type,\n'
    '  cbm: parseFloat(cbm),\n'
    '  origin_address: "출발지 주소",\n'
    '  port_name: port_name,\n'
    '  pickup_date: new Date().toISOString().split("T")[0],\n'
    '  mri_current: 0.72,\n'
    '  delay_reason: "해상 리스크 상승",\n'
    '  warehouse_name: warehouses[0].name,\n'
    '  warehouse_address: warehouses[0].address,\n'
    '  warehouse_km: warehouses[0].distance_km,\n'
    '  warehouse_minutes: warehouses[0].duration_min\n'
    '}\n'
    'On success: download the JSON response as a .json file\n'
    'Show success message: "루티 JSON 생성 완료!"'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. API 연결 코드
# ════════════════════════════════════════════════════════════════════════════
h1('7. API 연결 코드 작성법')
body('Lovable이 자동으로 코드를 만들어주지만, 직접 수정하거나 이해하고 싶다면 참고하세요.')

h2('7-1. API 설정 파일')
body('Lovable에 아래 파일을 만들어달라고 요청하세요:')
code(
    '// src/config/api.ts\n'
    'export const API_BASE = import.meta.env.VITE_API_BASE_URL\n'
    '                     ?? "http://localhost:8000";\n\n'
    'export async function apiGet(path: string) {\n'
    '  const res = await fetch(`${API_BASE}${path}`);\n'
    '  if (!res.ok) throw new Error(`API Error: ${res.status}`);\n'
    '  return res.json();\n'
    '}\n\n'
    'export async function apiPost(path: string, body: object) {\n'
    '  const res = await fetch(`${API_BASE}${path}`, {\n'
    '    method: "POST",\n'
    '    headers: { "Content-Type": "application/json" },\n'
    '    body: JSON.stringify(body),\n'
    '  });\n'
    '  if (!res.ok) throw new Error(`API Error: ${res.status}`);\n'
    '  return res.json();\n'
    '}'
)

h2('7-2. MRI 데이터 가져오기 예시')
code(
    '// MRI 데이터 불러오기\n'
    'import { useState, useEffect } from "react";\n'
    'import { apiGet } from "@/config/api";\n\n'
    'function Dashboard() {\n'
    '  const [mri, setMri] = useState(null);\n'
    '  const [loading, setLoading] = useState(true);\n\n'
    '  useEffect(() => {\n'
    '    apiGet("/api/mri")\n'
    '      .then(data => { setMri(data); setLoading(false); })\n'
    '      .catch(err => { console.error(err); setLoading(false); });\n'
    '  }, []);\n\n'
    '  if (loading) return <div>로딩 중...</div>;\n\n'
    '  return (\n'
    '    <div>\n'
    '      <h1>MRI: {mri?.mri}</h1>\n'
    '      <p>등급: {mri?.grade}</p>\n'
    '    </div>\n'
    '  );\n'
    '}'
)

h2('7-3. Lovable에게 코드 수정 요청하는 법')
body('코드가 잘못됐거나 수정이 필요하면 채팅창에 이렇게 말하세요:')
bullet('"API 호출이 실패할 때 에러 메시지를 한국어로 표시해줘"')
bullet('"MRI가 0.8 이상이면 배경을 빨간색으로 바꿔줘"')
bullet('"loading 중일 때 스피너를 가운데에 보여줘"')
tip('에러가 나면 브라우저 개발자 도구(F12) → Console 탭에서 에러 메시지를 복사해서 Lovable에 붙여넣으면 자동으로 수정해줍니다.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. 전체 테스트
# ════════════════════════════════════════════════════════════════════════════
h1('8. 전체 테스트 방법')

h2('8-1. 체크리스트')
table(
    ['번호', '테스트 항목', '기대 결과', '완료'],
    [
        ['1', '백엔드 서버 실행', 'http://localhost:8000/api/health → {"status":"ok"}', '☐'],
        ['2', '대시보드 로드', 'MRI 점수, 등급 배지 표시', '☐'],
        ['3', 'MRI 등급 색상', 'MRI에 따라 배지 색상 변경', '☐'],
        ['4', '과거 유사사례', '평균 지연일, 운임 상승률 표시', '☐'],
        ['5', 'LSTM 예측', '3개월 예측값 카드 표시', '☐'],
        ['6', '출하 등록 폼', '모든 필드 입력 가능', '☐'],
        ['7', '출하 제출', '결과 카드 (운임/지연/납기) 표시', '☐'],
        ['8', '창고 추천 연동', '창고 목록 3개 표시', '☐'],
        ['9', '4가지 옵션', 'A/B/C/D 비용 카드 표시, D에 별 표시', '☐'],
        ['10', 'JSON 생성 + 다운로드', '.json 파일이 다운로드됨', '☐'],
    ],
    col_widths=[1.2, 4, 5, 1.3]
)

h2('8-2. 테스트용 입력 데이터')
body('아래 값으로 전체 플로우를 테스트해보세요:')
code(
    '화주명: (주)경기식품\n'
    '화물 유형: 냉장화물\n'
    '화물 용량: 20 CBM\n'
    '희망 항로: 부산→로테르담\n'
    '집화 권역: 경기남부\n'
    '납기 여유: 14일\n'
    '긴급 화물: 체크 안 함'
)
body('예상 결과: MRI 0.6 이상이면 창고 추천 버튼이 활성화되고, '
     'D안 권장 총비용은 A안보다 낮게 나와야 합니다.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. 배포하기
# ════════════════════════════════════════════════════════════════════════════
h1('9. 배포하기')
body('개발이 완료되면 인터넷에서 누구나 접속할 수 있게 배포합니다.')

h2('9-1. 백엔드 배포 (Railway)')
body('Railway는 GitHub 연결만으로 Python 서버를 무료 배포할 수 있습니다.')

numberd('railway.app 접속 → "Start a New Project"', 1)
numberd('GitHub 계정 연결 → 프로젝트 저장소 선택', 2)
numberd('Deploy 설정:', 3)
code(
    'Root Directory: vs code  (api.py가 있는 폴더)\n'
    'Start Command:  uvicorn api:app --host 0.0.0.0 --port $PORT'
)
numberd('환경변수 설정: Railway 대시보드 → Variables 탭', 4)
code(
    'KAKAO_REST_API_KEY=실제키\n'
    'KAKAO_MOBILITY_KEY=실제키\n'
    'GEMINI_API_KEY=실제키\n'
    'ECOS_API_KEY=실제키'
)
numberd('배포 완료 후 제공되는 URL 복사 (예: https://wemet-api.railway.app)', 5)

tip('Railway 무료 플랜: 월 500시간 실행 가능. 데모/발표용으로 충분합니다.')

h2('9-2. 프론트엔드 배포 (Lovable 자동)')
body('Lovable은 작업 내용을 자동으로 배포해줍니다.')

numberd('Lovable 프로젝트 → Settings (설정)', 1)
numberd('Environment Variables 추가:', 2)
code('VITE_API_BASE_URL=https://wemet-api.railway.app')
numberd('저장 후 자동 재배포 완료', 3)
numberd('Lovable이 제공하는 URL로 접속 테스트', 4)

warn('배포 후 반드시 전체 테스트 체크리스트(8-1)를 다시 확인하세요.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 10. 트러블슈팅
# ════════════════════════════════════════════════════════════════════════════
h1('10. 트러블슈팅')

table(
    ['증상', '원인', '해결 방법'],
    [
        ['API 호출이 안 됨\n(CORS 에러)', '브라우저 보안 정책', '백엔드가 실행 중인지 확인\nhttp://localhost:8000/docs 접속 테스트'],
        ['창고 목록이 비어 있음', '카카오 API 키 없음', '정상 — 시뮬 DB 자동 사용\n키 입력 후 재시도시 실데이터'],
        ['MRI 점수가 0으로 나옴', '뉴스 수집 실패', '백엔드 터미널 에러 확인\nfeedparser 설치 여부 확인'],
        ['JSON 다운로드 안 됨', 'POST 요청 실패', '브라우저 F12 → Console 에러 확인\nLovable에 에러 메시지 붙여넣기'],
        ['화면이 깨져 보임', '스타일 문제', 'Lovable에 "Fix the layout on mobile" 요청'],
        ['배포 후 API 오류', '환경변수 미설정', 'Railway/Lovable Variables 설정 확인'],
    ],
    col_widths=[4, 3.5, 6]
)

h2('10-2. 도움 받는 방법')
bullet('백엔드 문제: api.py 코드 담당자에게 문의')
bullet('Lovable 코드 문제: F12 → Console 에러를 Lovable 채팅에 붙여넣기')
bullet('API 명세 확인: http://localhost:8000/docs (Swagger UI)')
bullet('전체 흐름 확인: 이 문서 4~6챕터 재확인')

divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('이 문서에 대한 문의는 백엔드 개발팀에 연락하세요.  |  '
          'API 문서: http://localhost:8000/docs').font.size = Pt(9)

# 저장
doc.save(OUT)
print(f'저장 완료: {OUT}')
