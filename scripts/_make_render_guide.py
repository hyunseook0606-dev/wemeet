# -*- coding: utf-8 -*-
"""Render 배포 가이드 DOCX 생성 — 친구용."""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print('python-docx 설치 필요: pip install python-docx')
    sys.exit(1)

OUT = Path(__file__).parent.parent / 'Render_배포_가이드_친구용.docx'
doc = Document()

s = doc.sections[0]
s.left_margin = s.right_margin = Cm(2.5)
s.top_margin  = s.bottom_margin = Cm(2.0)

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(10)

def numbered(text, num):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(10)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run(text)
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3CD')
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run('※ ' + text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x85, 0x64, 0x04)

def success(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D4EDDA')
    p._p.get_or_add_pPr().append(shd)
    r = p.add_run('✅ ' + text)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x15, 0x55, 0x24)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D6E4F0')
        hdr[i].paragraphs[0]._p.get_or_add_pPr().append(shd)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 표지
# ══════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
title_p.paragraph_format.space_after  = Pt(10)
tr = title_p.add_run('Render 배포 가이드')
tr.bold = True; tr.font.size = Pt(24)
tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(4)
sub_p.add_run('위밋모빌리티 해상 리스크 플랫폼 — FastAPI 백엔드 배포').font.size = Pt(12)

sub2_p = doc.add_paragraph()
sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2_p.paragraph_format.space_after = Pt(40)
sub2_p.add_run('친구용 (처음 해도 따라할 수 있습니다)').font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 전체 흐름 요약
# ══════════════════════════════════════════════════════════════
h1('전체 흐름 한눈에 보기')
body('아래 7단계를 순서대로 따라하면 됩니다. 총 소요 시간: 약 15~20분')
add_table(
    ['단계', '할 일', '소요 시간'],
    [
        ['1단계', 'GitHub 계정 만들기', '3분'],
        ['2단계', 'Render 가입하기', '2분'],
        ['3단계', 'Render에서 새 서비스 만들기', '3분'],
        ['4단계', '배포 설정 입력하기', '3분'],
        ['5단계', 'API 키 입력하기 (제일 중요!)', '3분'],
        ['6단계', '배포 시작 및 완료 확인', '5분'],
        ['7단계', 'Lovable에 서버 주소 연결', '2분'],
    ]
)
tip('팀원(본인)에게 GitHub 주소와 API 키를 미리 받아두세요. 5단계에서 필요합니다.')

# ══════════════════════════════════════════════════════════════
# 1단계
# ══════════════════════════════════════════════════════════════
h1('1단계 — GitHub 계정 만들기')
tip('이미 GitHub 계정이 있으면 이 단계는 건너뛰세요.')
numbered('브라우저(크롬 등)에서 github.com 접속', 1)
numbered('오른쪽 위 Sign up 클릭', 2)
numbered('이메일, 비밀번호, 아이디 입력 후 가입 완료', 3)

# ══════════════════════════════════════════════════════════════
# 2단계
# ══════════════════════════════════════════════════════════════
h1('2단계 — Render 가입하기')
numbered('브라우저에서 render.com 접속', 1)
numbered('Get Started for Free 클릭', 2)
numbered('GitHub으로 로그인 선택 (1단계에서 만든 계정으로)', 3)
numbered('GitHub 권한 허용 화면이 나오면 Authorize 클릭', 4)
success('Render 대시보드 화면이 나오면 성공입니다.')

# ══════════════════════════════════════════════════════════════
# 3단계
# ══════════════════════════════════════════════════════════════
h1('3단계 — 새 웹 서비스 만들기')
numbered('Render 대시보드에서 New + 버튼 클릭', 1)
numbered('Web Service 클릭', 2)
numbered('Connect a repository 화면에서 + Connect account 클릭 → GitHub 연결', 3)
numbered('팀원에게 받은 GitHub 주소의 저장소 선택', 4)
body('저장소 이름: wemeet-platform')
numbered('Connect 클릭', 5)
success('설정 입력 화면으로 넘어가면 성공입니다.')

# ══════════════════════════════════════════════════════════════
# 4단계
# ══════════════════════════════════════════════════════════════
h1('4단계 — 배포 설정 입력하기')
tip('아직 Create Web Service 버튼을 누르지 마세요. 5단계까지 마친 후 누릅니다.')
body('아래 표대로 입력하세요:')
add_table(
    ['항목', '입력값'],
    [
        ['Name', 'wemeet-api'],
        ['Region', 'Singapore'],
        ['Branch', 'main'],
        ['Build Command', 'pip install -r requirements-deploy.txt'],
        ['Start Command', 'uvicorn api:app --host 0.0.0.0 --port $PORT'],
        ['Instance Type', 'Free'],
    ]
)
tip('Build Command와 Start Command는 오타 없이 정확히 입력하세요. 복사 붙여넣기 권장.')

# ══════════════════════════════════════════════════════════════
# 5단계
# ══════════════════════════════════════════════════════════════
h1('5단계 — API 키 입력하기 (제일 중요!)')
body('같은 페이지를 아래로 스크롤하면 Environment Variables 섹션이 나옵니다.')
numbered('Add Environment Variable 클릭', 1)
numbered('아래 표의 Key와 Value를 하나씩 입력 (팀원에게 받은 키 사용)', 2)
add_table(
    ['Key (왼쪽 칸)', 'Value (오른쪽 칸)'],
    [
        ['GEMINI_API_KEY', '팀원에게 받은 Gemini 키'],
        ['KAKAO_REST_API_KEY', '팀원에게 받은 카카오 키'],
        ['KAKAO_MOBILITY_KEY', '카카오와 동일한 키'],
        ['ECOS_API_KEY', '팀원에게 받은 ECOS 키 (없으면 생략)'],
    ]
)
tip('키를 잘못 입력하면 창고 탐색, 보고서 기능이 작동하지 않습니다. 오타 주의!')

# ══════════════════════════════════════════════════════════════
# 6단계
# ══════════════════════════════════════════════════════════════
h1('6단계 — 배포 시작 및 완료 확인')
numbered('4단계, 5단계 입력이 모두 끝났으면 Create Web Service 버튼 클릭', 1)
numbered('배포 로그 화면이 나옵니다. 약 3~5분 기다립니다.', 2)
numbered('아래 메시지가 나오면 배포 성공:', 3)
code('INFO:     Application startup complete.')
code('INFO:     Uvicorn running on http://0.0.0.0:XXXXX')
numbered('페이지 상단에 서버 주소가 생성됩니다:', 4)
code('https://wemeet-api.onrender.com')
body('(주소는 서비스 이름에 따라 다를 수 있습니다)')
numbered('브라우저에서 아래 주소 접속해서 확인:', 5)
code('https://wemeet-api.onrender.com/api/health')
body('아래처럼 나오면 완전히 성공입니다:')
code('{"status": "ok", "timestamp": "2026-05-09T..."}')
success('이 주소를 복사해두세요. 7단계에서 사용합니다.')
tip('Render 무료 플랜은 15분 동안 사용하지 않으면 슬립 상태가 됩니다. 발표 전에 /api/health 주소를 한 번 접속해서 서버를 깨워두세요.')

# ══════════════════════════════════════════════════════════════
# 7단계
# ══════════════════════════════════════════════════════════════
h1('7단계 — Lovable에 서버 주소 연결하기')
numbered('Lovable 프로젝트 접속', 1)
numbered('오른쪽 위 Settings(설정) 클릭', 2)
numbered('Environment Variables 메뉴 클릭', 3)
numbered('아래 항목 추가 또는 수정:', 4)
add_table(
    ['Key', 'Value'],
    [
        ['VITE_API_BASE_URL', '6단계에서 복사한 Render 주소\n(예: https://wemeet-api.onrender.com)'],
    ]
)
tip('기존에 http://localhost:8000 으로 되어 있었다면 Render 주소로 교체하면 됩니다.')
numbered('저장 후 Lovable 프로젝트 새로고침', 5)
success('이제 Lovable 프론트엔드가 Render 서버와 연결됩니다. 어디서든 접속 가능!')

# ══════════════════════════════════════════════════════════════
# 자주 묻는 오류
# ══════════════════════════════════════════════════════════════
h1('자주 발생하는 오류와 해결 방법')
add_table(
    ['증상', '원인', '해결'],
    [
        ['배포 로그에 빨간 오류', 'Build Command 오타', 'Render 설정에서 Build Command 다시 확인'],
        ['/api/health 접속 안 됨', '아직 배포 중', '3~5분 더 기다린 후 새로고침'],
        ['창고 탐색 결과 0건', '카카오 키 오입력', 'Render 환경변수에서 KAKAO_REST_API_KEY 확인'],
        ['보고서 생성 안 됨', 'Gemini 키 오입력', 'Render 환경변수에서 GEMINI_API_KEY 확인'],
        ['Lovable에서 API 연결 안 됨', 'VITE_API_BASE_URL 미설정', '7단계 다시 확인'],
        ['첫 요청이 30초 걸림', 'Render 슬립 상태', '정상. 두 번째 요청부터 빠름. 발표 전 미리 깨울 것'],
    ]
)

# ══════════════════════════════════════════════════════════════
# 저장
# ══════════════════════════════════════════════════════════════
doc.save(OUT)
print(f'저장 완료: {OUT}')
